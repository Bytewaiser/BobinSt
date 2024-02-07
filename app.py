import streamlit as st
import pandas as pd
from PIL import Image

import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def img_to_doc(doc, name, img):
    # Add a heading and store the paragraph object
    heading = doc.add_heading(level=1)
    heading.add_run(f"{name}")

    # Center the heading
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the PIL image to a bytes buffer
    image_stream = io.BytesIO()
    img.save(image_stream, format="JPEG")  # Use 'PNG' if you prefer
    image_stream.seek(0)  # Go to the beginning of the stream

    image_paragraph = doc.add_paragraph()
    run = image_paragraph.add_run()
    run.add_picture(image_stream, width=Inches(4.0))

    # Center the paragraph, effectively centering the image
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


st.set_page_config(
    page_title="Buse Naming App",
)

st.title("Buse Naming App")


data = st.file_uploader(
    "Upload the excel file", type=["xlsx", "xls"], accept_multiple_files=False
)
# cols = [] if not data else pd.read_excel(data).columns
# index_col = st.selectbox("Select the index column", options=cols) # type: ignore
imgs = st.file_uploader(
    "Upload images", type=["jpg", "jpeg", "png"], accept_multiple_files=True
)


def get_data(selected):
    data = df[df["Sıralama"] == selected].squeeze()
    return data


if data:
    df = pd.read_excel(data)
    df.columns = df.columns.str.strip()
    index_col = "Sıralama"
    ids = df[index_col]
    st.write(df)

    selected_values = []  # type: ignore
    names = []
    images = []
    if imgs:
        tabs = st.tabs([f"Tab {idx+1}" for idx in range(len(imgs))])
        for idx, img in enumerate(imgs):
            image = Image.open(img)
            images.append(image)
            selected = tabs[idx].selectbox("Select the id", ids, key=idx)

            d = get_data(selected)

            machine = d["Çalışılan Makine"]
            no = d["Bobin No"]

            tabs[idx].image(image)
            selected_values.append(selected)
            name = f"{selected} NUMARALI BOBİN {machine} {no}"
            names.append(name)

        st.write(names)

if st.button("Create Docx"):
    doc = Document()
    for idx, name in enumerate(names):
        img_to_doc(doc, name, images[idx])

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    st.download_button("Download", data=doc_io, file_name="output.docx")
